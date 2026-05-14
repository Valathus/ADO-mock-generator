from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

# ── Styles helpers ──
def set_heading(paragraph, level=1):
    colors = {1: "1F3864", 2: "2E5E9E", 3: "17375E"}
    sizes  = {1: 18, 2: 14, 3: 12}
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.bold = True
    run.font.size = Pt(sizes.get(level, 12))
    run.font.color.rgb = RGBColor.from_string(colors.get(level, "000000"))

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    set_heading(p, 1)
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    set_heading(p, 2)
    return p

def h3(doc, text):
    p = doc.add_heading(text, level=3)
    set_heading(p, 3)
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(10.5)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.style.font.size = Pt(10.5)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F3864")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = cell_text
            if r_idx % 2 == 1:
                tc = row_cells[c_idx]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "EBF0FA")
                shd.set(qn("w:val"), "clear")
                tcPr.append(shd)
    return table

def code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F2F2")
    shading.set(qn("w:val"), "clear")
    p._p.get_or_add_pPr().append(shading)
    return p

def divider(doc):
    doc.add_paragraph("─" * 90)

# ═══════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("ADO Mock Generator")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor.from_string("1F3864")

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run("Technical Design Document (TDD)")
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor.from_string("2E5E9E")

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f"Version 1.0  •  {datetime.date.today().strftime('%B %d, %Y')}").font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. PURPOSE & SCOPE
# ═══════════════════════════════════════════════════════════════
h1(doc, "1. Purpose & Scope")
body(doc,
    "The ADO Mock Generator is a standalone Spring Boot application that simulates realistic "
    "Azure DevOps (ADO) activity against a live ADO Server instance. It is intended for "
    "demo environments, integration testing, and analytics pipeline validation where a "
    "consistent, controllable stream of ADO events (work items, repository pushes, pull "
    "requests, and pipeline builds) is required without involving real engineering teams."
)
doc.add_paragraph()
body(doc, "Primary objectives:")
bullet(doc, "Seed a configurable window of historical ADO data (work items, repo branches, build history) in one automated pass.")
bullet(doc, "Simulate ongoing sprint activity via a scheduled mutation engine that fires every 15 minutes.")
bullet(doc, "Remain fully idempotent and restart-safe: re-running against the same ADO project should converge rather than duplicate.")
bullet(doc, "Support a structured PI/Sprint iteration hierarchy matching typical SAFe-style ADO project layouts.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. SYSTEM OVERVIEW
# ═══════════════════════════════════════════════════════════════
h1(doc, "2. System Overview")

h2(doc, "2.1 High-Level Architecture")
body(doc,
    "The application is a headless Spring Boot process. It owns a single mutable JSON state "
    "file (`mock-state.json`) that tracks all seeded identifiers, sprint progress, and "
    "backfill status. Two Quartz scheduler jobs drive all activity:"
)
doc.add_paragraph()
add_table(doc,
    ["Job", "Trigger", "Responsibility"],
    [
        ["DataLoadMockJob", "Once on startup", "Initialize & validate state → seed iterations → seed work items → seed repo → seed builds"],
        ["WebhookMockJob",  "Cron: every 15 min (Mon–Sat, 05:00–23:00)", "Advance current sprint → expand sprint window → run mutation cycle (work items / repo / pipeline)"],
    ]
)

doc.add_paragraph()
h2(doc, "2.2 Technology Stack")
add_table(doc,
    ["Component", "Technology"],
    [
        ["Language / Runtime",  "Java 17"],
        ["Framework",           "Spring Boot 3.2.6"],
        ["HTTP Client",         "Spring WebFlux WebClient (Reactor Netty)"],
        ["Scheduler",           "Quartz 2.x (in-memory job store)"],
        ["JSON Serialization",  "Jackson 2 with JavaTimeModule"],
        ["Build Tool",          "Maven 3"],
        ["State Storage",       "Local filesystem JSON file (atomic file swap)"],
    ]
)

doc.add_paragraph()
h2(doc, "2.3 External Dependency")
body(doc,
    "The only external system is an Azure DevOps Server instance. The application communicates "
    "exclusively through the ADO REST API (version 6.0 by default) using HTTP Basic auth with "
    "Personal Access Tokens (PATs)."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. STATE MODEL
# ═══════════════════════════════════════════════════════════════
h1(doc, "3. State Model")

h2(doc, "3.1 MockState (mock-state.json)")
body(doc,
    "All runtime state is persisted in a single JSON file. The file is read fresh on every "
    "service call and written atomically (write-to-temp then rename) to prevent corruption."
)
doc.add_paragraph()
add_table(doc,
    ["Field", "Type", "Description"],
    [
        ["users",                    "List<User>",              "ADO user accounts (username + PAT + descriptor/id)"],
        ["admin",                    "Admin",                   "Admin account used for privileged operations"],
        ["collectionDetails",        "CollectionDetails",       "ADO URL, project name/id, team list, webhook URL"],
        ["dataLoadConfig",           "DataLoadConfig",          "Window days, sprint duration, items per day, builds per sprint, etc."],
        ["calendarAnchorDate",       "LocalDate",               "Fixed anchor; sprint 1 starts here. Set once, never changes."],
        ["currentSprintNumber",      "int",                     "Logical 'today's sprint'. Derived & advanced by IterationWebhookEngine."],
        ["iterationRootPath",        "String",                  "ADO path of the root Iterations node (e.g. \\MyProject\\Iteration)"],
        ["iterationRootIdentifier",  "String",                  "GUID of the root Iterations node"],
        ["programIterations",        "List<ProgramIteration>",  "PI + Sprint structure with ADO identifiers and work item / build IDs"],
        ["workItemBackfill",         "WorkItemBackfill",        "Backfill progress tracker (completed flag, last sprint, count in current sprint)"],
        ["repoBackfill",             "RepoBackfill",            "Repo seeding progress tracker"],
        ["buildBackfill",            "BuildBackfill",           "Build history seeding progress tracker"],
        ["repo",                     "RepoState",               "Repo ID, name, default branch head, pipeline definition ID"],
        ["webhookEnabled",           "boolean",                 "Master flag; mutation cycle will not run if false"],
        ["webhookMutationState",     "WebhookMutationState",    "Cycle counters and last-cycle timestamps for observability"],
    ]
)

doc.add_paragraph()
h2(doc, "3.2 Iteration Hierarchy")
body(doc,
    "The calendar model is fixed at startup. All dates are computed from calendarAnchorDate:"
)
doc.add_paragraph()
code_block(doc, "calendarAnchorDate = mostRecentMonday(today) - (sprintsPerPI × sprintDurationDays)\n\nPI N covers sprints: [(N-1)×sprintsPerPI+1 .. N×sprintsPerPI]\nSprint K starts at: anchorDate + (K-1) × sprintDurationDays\nSprint K ends  at: sprintStart + sprintDurationDays - 1")
body(doc,
    "PIs and Sprints are named using an abbreviated project label derived from the ADO project "
    "name (e.g. 'My Demo Project' → 'MDP-PI1', 'MDP-Sprint3')."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. COMPONENT DESIGN
# ═══════════════════════════════════════════════════════════════
h1(doc, "4. Component Design")

# 4.1
h2(doc, "4.1 FileStateRepository")
body(doc,
    "All reads and writes are synchronized. Writes use a write-to-temp-then-atomic-rename "
    "pattern. Jackson is configured with JavaTimeModule to correctly serialize/deserialize "
    "Instant and LocalDate fields."
)
bullet(doc, "load() — deserializes JSON file; returns new MockState() if file is absent.")
bullet(doc, "save(state) — serializes to .tmp file then atomically moves it to the state path.")

doc.add_paragraph()

# 4.2
h2(doc, "4.2 AdoRestClient")
body(doc,
    "Thin wrapper around WebClient. A new WebClient is created per call (reads URL + PAT from "
    "state at call time to avoid stale values). Supports: GET, POST, PATCH, PUT, postJsonPatch, "
    "patchJsonPatch. All methods apply a configurable fixed-delay retry (default 3 attempts, "
    "500 ms backoff). Error bodies are surfaced as RuntimeException messages."
)

doc.add_paragraph()

# 4.3
h2(doc, "4.3 Identity Providers")
add_table(doc,
    ["Class", "Behaviour"],
    [
        ["AdminIdentityProvider", "Always returns state.admin; throws if PAT is missing."],
        ["UserIdentityProvider",  "Caches enabled users on @PostConstruct. next() does round-robin cycling. random() picks randomly. allUsers() returns the full cached list."],
    ]
)

doc.add_paragraph()

# 4.4
h2(doc, "4.4 MockStateInitializer")
body(doc,
    "Called at the start of DataLoadMockJob. Validates required fields (users, admin, "
    "collectionDetails, dataLoadConfig) and normalizes zero/negative config values to safe "
    "defaults (e.g. windowDays → 365, sprintDurationDays → 14). Returns false to abort the "
    "job if validation fails."
)

doc.add_paragraph()

# 4.5
h2(doc, "4.5 IterationService")
body(doc,
    "Responsible for creating and maintaining the PI/Sprint iteration tree inside ADO."
)
bullet(doc, "ensureSprintStructure() — one-time bootstrap. Computes anchor date, calls ensureSprintsExist(1..N), sets iterationSeedCompleted=true.")
bullet(doc, "ensureSprintsExist(start, end) — idempotent range creator. Loads the ADO iteration tree once, collects existing paths into a HashSet, then creates only missing PIs and Sprints via POST. Assigns new sprints to all configured teams.")
bullet(doc, "createIfMissing() — normalizes paths for reliable existence checks; fetches identifier via GET if already present.")
body(doc, "Naming follows IterationNamingSupport: project abbreviation + \"-PI{n}\" / \"-Sprint{n}\".")

doc.add_paragraph()

# 4.6
h2(doc, "4.6 WorkItemsCreationService")
body(doc,
    "Two modes of operation:"
)
bullet(doc, "isWorkItemsDataLoaded() — historical backfill. Iterates through all sprints in the window, creates (workItemsPerDay × sprintDurationDays) items per sprint. Uses a batch-save mechanism (every 25 items by default) to limit data loss on crash. Resumes from the last processed sprint/index on restart.")
bullet(doc, "ensureWorkItemsForSprint(state, sprint, correlation) — webhook delta. Idempotent: creates only the delta between sprint.workItemIds.size() and the expected count.")
body(doc,
    "Work item type distribution per sprint (index mod 10): index 0 → Requirement type, "
    "index 1 → Bug type, indices 2–9 → Task type. Subtask selection for PBIs is deterministic "
    "per sprint using a stable CRC32-seeded LCG, ensuring the same items get subtasks on "
    "restart."
)

doc.add_paragraph()

# 4.7
h2(doc, "4.7 RepoSeederService")
body(doc, "Bootstrap seeder for the Git repository:")
bullet(doc, "Resolves or creates the repo by name via the ADO Git Repositories API.")
bullet(doc, "Ensures the default branch exists by pushing an initial commit (README.md) if the ref is absent.")
bullet(doc, "Creates environment branches (dev, qa, demo, prod) from the default branch head using the refs API (not an empty commit — avoids ADO rejection).")
bullet(doc, "ensureRepoActivityForSprint() creates a feature/sprint-N-1 branch off dev for each new sprint during webhook expansion.")

doc.add_paragraph()

# 4.8
h2(doc, "4.8 PipelineSeederService")
body(doc, "Two modes:")
bullet(doc, "seedBuildHistoryIfRequired() — one-time seed: queues buildsPerSprint builds against the default branch with deterministic scenario rotation (FAIL / SUCCESS / RETRY / CANCEL / SUCCESS on index mod 5).")
bullet(doc, "ensureBuildsForSprint() — webhook delta: idempotent, adds only missing builds up to the expected count per sprint.")
body(doc, "Resolves the build definition ID lazily by matching pipeLineName against the ADO build definitions list.")

doc.add_paragraph()

# 4.9
h2(doc, "4.9 MutationExecutor")
body(doc,
    "Orchestrator for the webhook simulation cycle. Called by WebhookMockJob."
)
bullet(doc, "Guards: skips if state.webhookEnabled=false or work item backfill is not yet complete.")
bullet(doc, "Calls iterationWebhookEngine.tickAndMaybeSeed() first to advance the current sprint and expand the window if needed.")
bullet(doc, "Dispatches mutationsPerHour mutations per cycle. Each mutation is randomly routed:")
doc.add_paragraph()
add_table(doc,
    ["% (default)", "Action", "Engine"],
    [
        ["50%", "Update / create a work item", "WorkItemWebhookEngine"],
        ["30%", "Run a repo workflow cycle (push / PR / merge)", "RepoWebhookEngine"],
        ["20%", "Queue or cancel a pipeline build", "PipelineWebhookEngine"],
    ]
)
body(doc, "Each action receives a unique correlation string (mockRun:batchId:N) for tracing through ADO audit logs and tags.")

doc.add_paragraph()

# 4.10
h2(doc, "4.10 WorkItemWebhookEngine")
body(doc, "Mutation mix (applied to each work item action):")
add_table(doc,
    ["% (of update actions)", "Mutation"],
    [
        ["20% (configurable)", "CREATE a new Task in the current sprint iteration"],
        ["35%", "Change state (weighted realism: New→Active 70%, Active→Closed 60%, etc.)"],
        ["20%", "Change assignee to a random enabled user"],
        ["15%", "Carry-over to next sprint (only fires at sprint-close boundary, 10% of the time)"],
        ["10%", "Add comment to history"],
        ["10%", "Append tag (backend / frontend / hotfix) — reads existing tags first to avoid overwrite"],
        ["10%", "Add a related link to another work item in the current sprint"],
    ]
)

doc.add_paragraph()

# 4.11
h2(doc, "4.11 RepoWebhookEngine")
body(doc, "Simulates a corporate Git-flow pattern:")
add_table(doc,
    ["% per cycle", "Action"],
    [
        ["1%",  "Create a new feature branch from dev HEAD + open a PR"],
        ["24%", "Open a PR from an existing feature branch"],
        ["25%", "Push a new commit to a random feature branch"],
        ["50%", "Mutate a random active PR (vote / comment / add reviewer / push / complete / abandon)"],
        ["60%", "Promote dev → qa (probabilistic, independent roll)"],
        ["40%", "Promote qa → demo (probabilistic, independent roll)"],
        ["25%", "Promote demo → prod (probabilistic, independent roll)"],
    ]
)
body(doc, "Work items are linked to PRs via ArtifactLink (vstfs:///Git/PullRequestId/...) on creation.")

doc.add_paragraph()

# 4.12
h2(doc, "4.12 PipelineWebhookEngine")
body(doc, "Per webhook tick, resolves a scenario (SUCCESS 50% / FAIL 30% / CANCEL 20%) and:")
bullet(doc, "SUCCESS or FAIL → queues a new build against a randomly chosen branch (dev 40% / qa 30% / demo 20% / prod 10%).")
bullet(doc, "CANCEL → finds a currently in-progress build and patches its status to 'cancelling'.")

doc.add_paragraph()

# 4.13
h2(doc, "4.13 IterationWebhookEngine")
body(doc,
    "Called once per webhook cycle before mutation. Computes the expected current sprint "
    "number from calendarAnchorDate and today's date, advances state.currentSprintNumber "
    "if behind (forward-only), then checks whether the sprint window needs expanding "
    "(today + windowDays). For each newly required sprint, calls:"
)
bullet(doc, "iterationService.ensureSprintsExist() — creates the ADO iteration nodes")
bullet(doc, "workItemsCreationService.ensureWorkItemsForSprint() — seeds work items")
bullet(doc, "repoSeederService.ensureRepoActivityForSprint() — seeds a feature branch")
bullet(doc, "pipelineSeederService.ensureBuildsForSprint() — seeds builds")

doc.add_paragraph()

# 4.14
h2(doc, "4.14 WebhookServiceManager")
body(doc,
    "Ensures ADO service hook subscriptions exist for: workitem.created, workitem.updated, "
    "git.push, git.pullrequest.created, git.pullrequest.updated, build.complete. "
    "Subscriptions are checked by eventType + webhook URL to avoid duplicates. Only runs "
    "when state.webhookEnabled=true."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. DATA FLOW
# ═══════════════════════════════════════════════════════════════
h1(doc, "5. Data Flow")

h2(doc, "5.1 Initial Data Load Sequence")
code_block(doc,
"""Application startup
  └─ DataLoadMockJob.execute()
       ├─ MockStateInitializer.initializeAndValidate()   → validate & normalize mock-state.json
       ├─ IterationService.ensureSprintStructure()
       │    ├─ Compute calendarAnchorDate (mostRecentMonday - pastDays)
       │    ├─ Call ADO: GET /project/_apis/wit/classificationnodes/iterations
       │    └─ For each sprint in [1..requiredSprintCount]:
       │         ├─ POST PI node if missing
       │         ├─ POST Sprint node if missing
       │         └─ POST sprint to each team
       ├─ WorkItemsCreationService.isWorkItemsDataLoaded()
       │    └─ For each sprint:
       │         └─ For each item index [0..workItemsPerDay×sprintDuration):
       │              ├─ POST work item (with bypassRules=true)
       │              └─ POST subtask children (deterministic set)
       ├─ RepoBackfillController.runBatchIfRequired()
       │    ├─ RepoSeederService.seedIfRequired()
       │    │    ├─ GET/POST repo
       │    │    ├─ Ensure default branch (initial commit)
       │    │    └─ Ensure env branches (dev/qa/demo/prod via refs API)
       │    └─ RepoWorkflowEngine.runWorkflowCycle() × repoActivitiesPerSprint
       │         └─ feature branch → PR → complete
       └─ PipelineSeederService.seedBuildHistoryIfRequired()
            └─ Queue buildsPerSprint builds with scenario rotation"""
)

doc.add_paragraph()

h2(doc, "5.2 Webhook Mutation Cycle")
code_block(doc,
"""WebhookMockJob.execute()  [every 15 min]
  └─ (guard: state.webhookEnabled == true)
  └─ MutationExecutor.runCycle()
       ├─ IterationWebhookEngine.tickAndMaybeSeed(state)
       │    ├─ Advance currentSprintNumber if calendar moved forward
       │    └─ For each new sprint [existingMax+1 .. requiredMax]:
       │         ├─ IterationService.ensureSprintsExist()
       │         ├─ WorkItemsCreationService.ensureWorkItemsForSprint()
       │         ├─ RepoSeederService.ensureRepoActivityForSprint()
       │         └─ PipelineSeederService.ensureBuildsForSprint()
       └─ For i in [1..mutationsPerHour]:
            ├─ roll < 50% → WorkItemWebhookEngine.runWebhookUpdate()
            ├─ roll < 80% → RepoWebhookEngine.runWebhookCycle()
            └─ roll ≥ 80% → PipelineWebhookEngine.runPipelineWebhook()"""
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. CONFIGURATION REFERENCE
# ═══════════════════════════════════════════════════════════════
h1(doc, "6. Configuration Reference")

h2(doc, "6.1 application.properties Keys")
add_table(doc,
    ["Property", "Default", "Description"],
    [
        ["ado.apiVersion",                        "6.0",    "ADO REST API version"],
        ["ado.connectTimeoutMs",                  "5000",   "HTTP connect timeout (ms)"],
        ["ado.readTimeoutMs",                     "15000",  "HTTP read timeout (ms)"],
        ["ado.retry.maxAttempts",                 "3",      "Retry count on HTTP error"],
        ["ado.retry.backoffMs",                   "500",    "Fixed backoff between retries (ms)"],
        ["seed.windowDays",                       "7",      "Historical window for seeding"],
        ["seed.workItemsPerDay",                  "2",      "Work items created per calendar day"],
        ["seed.buildsPerDay",                     "5",      "Builds queued per day (used in seeder)"],
        ["mock.seed.workItems.enabled",           "true",   "Enable work item seeding"],
        ["mock.seed.builds.enabled",              "true",   "Enable build seeding"],
        ["mock.seed.repo.enabled",                "true",   "Enable repo seeding"],
        ["mock.seed.iterations.enabled",          "true",   "Enable iteration seeding"],
        ["mock.mutation.enabled",                 "false",  "Enable mutation cycle (overridden by state.webhookEnabled)"],
        ["mock.webhooks.enabled",                 "false",  "Enable ADO webhook subscription creation"],
        ["mock.random.seed",                      "12345",  "Base seed for deterministic mutation randomness"],
        ["mutations.perHour",                     "10",     "Mutations dispatched per webhook cycle"],
        ["mutations.mix.workItemUpdatePct",       "50",     "% of mutations that update work items"],
        ["mutations.mix.repoWorkflowPct",         "30",     "% of mutations that run repo workflows"],
        ["mutations.mix.pipelinePct",             "20",     "% of mutations that trigger pipelines"],
        ["mock.pipeline.cancel.enabled",          "true",   "Allow pipeline cancel scenario"],
        ["mock.pipeline.updateWorkItemOnSuccess", "true",   "Update work item on build success"],
        ["mock.buildQueueName",                   "Default","Build queue/pool name"],
        ["storage.stateFile",                     "./mock-state.json", "Path to the JSON state file"],
    ]
)

doc.add_paragraph()

h2(doc, "6.2 mock-state.json Required Fields (bootstrap)")
body(doc, "Before the first run, the following fields must be populated manually in mock-state.json:")
add_table(doc,
    ["Field path", "Example", "Notes"],
    [
        ["admin.username",                    "adoadmin",                       "ADO admin username"],
        ["admin.pat",                         "<token>",                        "Admin Personal Access Token"],
        ["users[].username",                  "dev1@contoso.com",               "At least one user required"],
        ["users[].pat",                       "<token>",                        "User PAT"],
        ["collectionDetails.url",             "https://ado.contoso.com/tfs",    "ADO Server collection URL"],
        ["collectionDetails.projectName",     "My Demo Project",                "Target ADO project"],
        ["collectionDetails.projectId",       "<guid>",                         "ADO project GUID"],
        ["collectionDetails.webhookURL",      "https://receiver.example.com",   "URL for ADO service hook subscriptions"],
        ["collectionDetails.teams",           "[\"Team Alpha\"]",               "Teams to assign sprints to (can be empty)"],
        ["repo.repoName",                     "mock-repo",                      "Git repo name to seed"],
        ["repo.defaultBranch",                "main",                           "Default branch name"],
        ["repo.pipeLineName",                 "mock-build",                     "Build definition name for pipeline seeding"],
        ["dataLoadConfig.windowDays",         "90",                             "Total calendar window to seed"],
        ["dataLoadConfig.sprintDurationDays", "14",                             "Sprint length in days"],
        ["dataLoadConfig.sprintsPerPI",       "4",                              "Sprints per Program Increment"],
        ["dataLoadConfig.workItemsPerDay",    "3",                              "Work items per calendar day per sprint"],
        ["dataLoadConfig.buildsPerSprint",    "2",                              "Builds to seed per sprint"],
        ["dataLoadConfig.repoActivitiesPerSprint", "3",                        "Repo workflow cycles during initial seeding"],
        ["dataLoadConfig.currentSprintFillPercentage", "70",                   "Fill percentage for in-progress sprint"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. IDEMPOTENCY & RESTART SAFETY
# ═══════════════════════════════════════════════════════════════
h1(doc, "7. Idempotency & Restart Safety")

body(doc,
    "The entire seeding pipeline is designed to be crash-safe. The following mechanisms "
    "enforce idempotency:"
)
doc.add_paragraph()
add_table(doc,
    ["Concern", "Mechanism"],
    [
        ["Iteration tree",        "Existence check via full path set before any POST. Already-existing nodes are fetched (GET identifier) rather than re-created."],
        ["Work item backfill",    "state.workItemBackfill.lastProcessedSprintNumber + createdInCurrentSprint tracks exact resume position. Batch-saves every N items."],
        ["Repo bootstrap",        "getBranchHeadObjectId() check before any branch creation. Returns early if ref already exists."],
        ["Build seeding",         "sprint.buildIds.size() vs expected count guard. Only the delta is queued."],
        ["Webhook subscriptions", "Subscription list is fetched and compared by eventType + URL before creating."],
        ["Work items (webhook)",  "sprint.workItemIds.size() >= expected → no-op."],
        ["PI/Sprint (webhook)",   "ensureSprintsExist() reuses the same existence-check logic as the initial seed."],
        ["State file write",      "Atomic rename (write to .tmp, then move). Falls back to non-atomic copy if OS does not support atomic move cross-device."],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 8. SCHEDULING DESIGN
# ═══════════════════════════════════════════════════════════════
h1(doc, "8. Scheduling Design")

add_table(doc,
    ["Job", "Identity", "Trigger type", "Schedule", "Concurrency"],
    [
        ["DataLoadMockJob", "dataLoadMockJob", "Simple (fire once)", "Immediately on startup, no repeat", "@DisallowConcurrentExecution"],
        ["WebhookMockJob",  "webhookMockJob",  "Cron",               "0 0/15 5-23 ? * MON-SAT (every 15 min, working hours)", "@DisallowConcurrentExecution"],
    ]
)

doc.add_paragraph()
body(doc,
    "Both jobs are annotated with @DisallowConcurrentExecution to prevent overlapping "
    "executions. The Quartz job store is in-memory, so schedule state is not persisted across "
    "restarts. The data load job will re-run on each startup but will exit immediately if "
    "iterationSeedCompleted, workItemBackfill.completed, buildBackfill.completed, and "
    "repoBackfill.completed are all true in state."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 9. ADO API USAGE SUMMARY
# ═══════════════════════════════════════════════════════════════
h1(doc, "9. ADO API Usage Summary")

add_table(doc,
    ["Domain", "API endpoint pattern", "Method", "Purpose"],
    [
        ["Iterations", "/{project}/_apis/wit/classificationnodes/iterations", "GET / POST", "Read tree / create PI or Sprint node"],
        ["Iterations (team)", "/{project}/{team}/_apis/work/teamsettings/iterations", "POST", "Assign sprint to team"],
        ["Work Items", "/{project}/_apis/wit/workitems/${type}", "POST (json-patch+json)", "Create work item"],
        ["Work Items", "/{project}/_apis/wit/workitems/{id}", "PATCH (json-patch+json)", "Update work item fields"],
        ["Work Item Types", "/{project}/_apis/wit/workitemtypes", "GET", "Resolve type names by category"],
        ["Repositories", "/{project}/_apis/git/repositories", "GET / POST", "Resolve or create repo"],
        ["Git Pushes", "/{project}/_apis/git/repositories/{id}/pushes", "POST", "Create branch / push commit"],
        ["Git Refs", "/{project}/_apis/git/repositories/{id}/refs", "GET / POST", "Read branch heads / create branch ref"],
        ["Pull Requests", "/{project}/_apis/git/repositories/{id}/pullrequests", "GET / POST / PATCH", "Create, query, update PRs"],
        ["PR Reviewers", "/{project}/_apis/git/repositories/{id}/pullRequests/{prId}/reviewers/{uid}", "PUT", "Add / update reviewer vote"],
        ["PR Threads", "/{project}/_apis/git/repositories/{id}/pullRequests/{prId}/threads", "POST", "Add comment thread"],
        ["PR Work Items", "/{project}/_apis/git/repositories/{id}/pullRequests/{prId}/workitems", "GET", "Check linked work items"],
        ["Builds", "/{project}/_apis/build/builds", "GET / POST / PATCH", "Queue build / query / cancel"],
        ["Build Definitions", "/{project}/_apis/build/definitions", "GET", "Resolve definition ID by name"],
        ["Identities", "/_apis/Identities", "GET", "Resolve user descriptor/GUID by username"],
        ["Webhooks", "/_apis/hooks/subscriptions", "GET / POST", "List / create service hook subscriptions"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 10. SECURITY CONSIDERATIONS
# ═══════════════════════════════════════════════════════════════
h1(doc, "10. Security Considerations")

add_table(doc,
    ["Concern", "Current approach", "Recommendation"],
    [
        ["PAT storage",         "PATs stored in plaintext in mock-state.json on disk",       "Use OS secret store or environment variable injection. Restrict file permissions to the service account."],
        ["PAT rotation",        "No automated rotation; PAT expiry will cause hard failures", "Monitor ADO 401 responses and alert. Consider short-lived PATs with renewal scripting."],
        ["Network",             "HTTP Basic auth over TLS to ADO Server",                    "Verify TLS certificate chain. Avoid self-signed certs in production environments."],
        ["bypassRules flag",    "Used for work item creation to set backdated CreatedDate",   "Limit admin PAT scope. Do not use bypassRules in environments where audit integrity matters."],
        ["suppressNotifications", "Used on all work item creates to avoid email spam",        "Acceptable for demo environments. Disable if notification testing is required."],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 11. KNOWN LIMITATIONS & FUTURE WORK
# ═══════════════════════════════════════════════════════════════
h1(doc, "11. Known Limitations & Future Work")

h2(doc, "11.1 Current Limitations")
bullet(doc, "No HTTP server / REST API exposed. The application cannot be queried or controlled at runtime without editing mock-state.json and restarting.")
bullet(doc, "Work item type resolution is done once per process lifetime (workItemTypesLoaded flag). A process restart is required if the ADO project's work item types change.")
bullet(doc, "UserIdentityProvider caches users at startup. Adding users to mock-state.json requires a restart to take effect.")
bullet(doc, "RepoWorkflowEngine.runWorkflowCycle() always uses users[0] as the author for backfill commits, reducing author diversity in historical data.")
bullet(doc, "The WebhookMockJob scheduler does not persist misfire history — if the application is down during a scheduled window, those cycles are skipped.")
bullet(doc, "No dead-letter or error-retry queue for failed ADO API calls within a mutation cycle; failed actions are counted and logged but not retried in the same cycle.")
bullet(doc, "Iteration naming abbreviation logic truncates to 3 characters, which may produce collisions for projects with similar names.")

doc.add_paragraph()
h2(doc, "11.2 Potential Improvements")
bullet(doc, "Expose a Spring MVC/WebFlux REST endpoint to trigger cycles, inspect state, and toggle feature flags without a restart.")
bullet(doc, "Replace file-based state with an embedded database (H2/SQLite) for concurrent-safe multi-process support.")
bullet(doc, "Add Spring Boot Actuator health/metrics endpoints for observability.")
bullet(doc, "Support environment-variable-driven PAT injection to avoid secrets in the JSON state file.")
bullet(doc, "Add integration tests using WireMock to simulate ADO API responses.")
bullet(doc, "Extend PipelineWebhookEngine to update work item status on build success/failure (currently flagged by properties but not fully wired).")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 12. GLOSSARY
# ═══════════════════════════════════════════════════════════════
h1(doc, "12. Glossary")

add_table(doc,
    ["Term", "Definition"],
    [
        ["ADO",             "Azure DevOps Server — the Microsoft ALM platform being targeted"],
        ["PAT",             "Personal Access Token — credential used for ADO API authentication"],
        ["PI",              "Program Increment — a planning cycle containing multiple sprints (SAFe terminology)"],
        ["Anchor Date",     "The fixed Monday from which all sprint dates are computed"],
        ["Backfill",        "One-time seeding of historical data within a configured date window"],
        ["Mutation Cycle",  "One execution of MutationExecutor.runCycle() — dispatches N mutations across work items, repos, and pipelines"],
        ["Correlation ID",  "A unique string (mockRun:batchId:N) injected into tags/history for end-to-end traceability"],
        ["bypassRules",     "ADO API flag that allows setting system fields (e.g. CreatedDate) that are otherwise read-only"],
        ["ArtifactLink",    "A vstfs:// URL that links a work item to a pull request in ADO's relationship model"],
        ["Idempotent",      "An operation that produces the same result whether called once or many times"],
    ]
)

doc.add_paragraph()

# ── Footer note ──
doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}  •  ado-mock-generator v1.0.0")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(150, 150, 150)

# ── Save ──
out_path = "/Users/srimanK/Downloads/ado-mock-generator/ADO_Mock_Generator_TDD.docx"
doc.save(out_path)
print(f"Saved: {out_path}")